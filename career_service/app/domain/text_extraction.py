from abc import ABC, abstractmethod
import pdfplumber
from docx import Document
import logging
from dataclasses import dataclass
from typing import Dict, List
from pathlib import Path
import re

from service.service_file import get_file_extension

# Configure logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class TextExtractionResult:
    filename: str
    total_pages: int
    pages: Dict[int, str]  # page_number -> text
    metadata: Dict[str, str]

class FileManager(ABC):
    @abstractmethod
    def extract_text(self, file_path: str) -> TextExtractionResult:
        """Extract text from file and return structured result"""
        pass

class PDFFileManager(FileManager):
    def extract_text(self, file_path: str) -> TextExtractionResult:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = {}
                for i, page in enumerate(pdf.pages, 1):
                    raw_text = page.extract_text() or ""
                    # Replace tabs with space
                    raw_text = raw_text.replace('\t', ' ')
                    # Remove page break characters
                    raw_text = raw_text.replace('\x0C', '')
                    # Replace special characters with space
                    raw_text = re.sub(r'[^\x20-\x7E]+', ' ', raw_text)
                    # Replace multiple consecutive spaces with just one
                    raw_text = re.sub(r'\s+', ' ', raw_text).strip()

                    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
                    cleaned_text = "\n".join(lines)
                    pages[i] = cleaned_text

                return TextExtractionResult(
                    filename=Path(file_path).name,
                    total_pages=len(pdf.pages),
                    pages=pages,
                    metadata={
                        "producer": pdf.metadata.get("Producer", ""),
                        "creator": pdf.metadata.get("Creator", "")
                    }
                )
        except Exception as e:
            return TextExtractionResult(
                filename=Path(file_path).name,
                total_pages=0,
                pages={},
                metadata={"error": str(e)}
            )

class DOCXFileManager(FileManager):
    def extract_text(self, file_path: str) -> TextExtractionResult:
        try:
            doc = Document(file_path)
            
            # Collect non-blank paragraphs, removing page break characters
            paragraphs = []
            for p in doc.paragraphs:
                text = p.text.replace('\x0C', '').strip()
                if text:
                    paragraphs.append(text)

            # Collect non-blank table cells, removing page break characters
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.replace('\x0C', '').strip()
                        if cell_text:
                            tables.append(cell_text)

            # Join lines with a newline
            combined_text = "\n".join(paragraphs + tables)

            return TextExtractionResult(
                filename=Path(file_path).name,
                total_pages=1,
                pages={1: combined_text},
                metadata={}
            )
        except Exception as e:
            return TextExtractionResult(
                filename=Path(file_path).name,
                total_pages=0,
                pages={},
                metadata={"error": str(e)}
            )

class TXTFileManager(FileManager):
    def extract_text(self, file_path: str) -> TextExtractionResult:
        try:
            with open(file_path, 'r') as file:
                lines = [line for line in file.read().splitlines() if line.strip()]
                content = "\n".join(lines)
                return TextExtractionResult(
                    filename=Path(file_path).name,
                    total_pages=1,
                    pages={1: content},
                    metadata={}
                )
        except Exception as e:
            return TextExtractionResult(
                filename=Path(file_path).name,
                total_pages=0,
                pages={},
                metadata={"error": str(e)}
            )

# Usage example
def process_document(file_path: str, file_manager: FileManager) -> TextExtractionResult:
    return file_manager.extract_text(file_path)

def get_text_from_file(file_info):
    logger.info(f"Extracting text from file: {file_info.filename} and whole data from {file_info}")
    extension = get_file_extension(file_info.filename)
    if extension == "pdf":
        manager = PDFFileManager()
    elif extension == "docx":
        manager = DOCXFileManager()
    elif extension == "txt":
        manager = TXTFileManager()
    else:
        raise ValueError(f'Unsupported file type {extension}')

    logger.info(f"Using {manager.__class__.__name__} to extract text")
    return manager.extract_text(file_info.save_path)